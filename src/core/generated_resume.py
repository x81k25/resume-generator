# external imports
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from openai import OpenAI
import pickle
import time
# internal imports
from src.utils.single_content_completion import complete_single_content
from src.utils.logger import log

# ------------------------------------------------------------------------------
# define primary class
# ------------------------------------------------------------------------------

class GeneratedResume:
    def __init__(
        self,
        env_vars,
        model_config,
        doc_format,
        job_description,
        resume_input,
        role_title_overrides = None
    ):
        log("initializing GeneratedResume object")
        # ingested file parameters
        self.env_vars = env_vars
        self.model_config = model_config
        self.doc_format = doc_format
        self.job_description = job_description
        self._set_gen_resume_components(resume_input)
        # input parameters
        # will overwrite role titles given for each employer
        if role_title_overrides is not None:
            self.role_title_overrides = role_title_overrides
        else:
            self.role_title_overrides = [None] * len(self.professional_experience_input)
        # # elements to be generated internally
        self.gen_soft_skills = None
        self.gen_tech_tools = None
        self.gen_tech_skills = None
        self.professional_experience_output = []  # will hold the final output to be given to resume writer
        log("GeneratedResume object initialized")

    def _set_gen_resume_components(self, resume_input):
        self.professional_experience_liminal = [] # will hold the intermediate data in processing created for the professional_experience_output
        self.personal_info = resume_input['personal_info']
        self.professional_experience_input = resume_input['professional_experience']
        self.education = resume_input['education']
        self.military_experience = resume_input['military_experience']
        self.hard_skills = resume_input['hard_skills']
        for experience in self.professional_experience_input:
            self.professional_experience_liminal.append({"employer": experience["employer"]})

# ------------------------------------------------------------------------------
# helper functions to be used by other functions in this class
# ------------------------------------------------------------------------------

    @staticmethod
    def _parse_response(response):
        """
        parses the response from the OpenAI chat completion API
        :param response: response from the API
        :return: the json dict version of the response
        """
        try:
            output = response.lstrip("```json\n").rstrip("\n```")
            output = json.loads(output)
        except:
            raise ValueError(f"Error: \"{response}\" is not the correct format for parse_response")

        return output


    @staticmethod
    def _de_parse_response(response):
        """
        revert the parsing operation to normalize the text for OpenAI
        :param response: response from the API
        :return: the json dict version of the response
        """
        output = "```json\n" + str(response) + "\n```"
        return output


    @staticmethod
    def _generate_role_title(client, responsibilities):
        """
        calls the OpenAI chat completion API and returns the title of the role
        :param client: client API object
        :param responsibilities: list of responsibilities for 1 role
        :return: the tile of the role
        """
        role_title = complete_single_content(
            client,
            "Return only one job title given the following list" +
            responsibilities
        )

        return role_title

    @staticmethod
    def add_line(doc):
        """
        Add a horizontal line to the document
        :param doc:
        :return:
        """
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(1)
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture('./data/assets/black-line.png', width=Inches(6.375))
        font = run.font
        font.size = Pt(1)
        p_format = p.paragraph_format
        p_format.left_indent = Inches(-0.03125)


# ------------------------------------------------------------------------------
# sub-functions over the over-arching generate_resume function below
#
# these functions will generally be called together, but due to API consumption
# need to be externally callable for testing
# ------------------------------------------------------------------------------

    def generate_resume_content(self):
        """
         Generates a resume based on a job description and a list of experiences
         :return:
         """
        log("generating resume content")

        # instantiate client
        client = OpenAI()

        # extract key technical skills from the job description
        self.gen_tech_skills = complete_single_content(
            client,
            "Extract the technical skills " +
            "required within this job description: " + self.job_description['role_description'] + " " +
            "Do not return any other additional skills",
            self.model_config['resume_gen_temp']
        )

        # extract key technical tools from the job description
        self.gen_tech_tools = complete_single_content(
            client,
            "Extract all technology tools, e.g. coding languages, cloud " +
            "development tools, and any specific development methodologies  " +
            "required within this job description: " + self.job_description['role_description'] + " " +
            "Do not return any other additional skills",
            self.model_config['resume_gen_temp']
        )

        # extract soft skills from the job description
        self.gen_soft_skills = complete_single_content(
            client,
            "Extract the key soft skills from this job description:" + self.job_description['role_description'],
            self.model_config['resume_gen_temp']
        )

        # select key experience from professional experience for each employer
        for i in range(len(self.professional_experience_input)):
            response = complete_single_content(
                client,
                "From this list: " + str(self.professional_experience_input[i]['responsibilities']) + " " +
                "select the JSON elements that correspond to the following skills: " +
                self.gen_tech_skills + self.gen_tech_tools + self.gen_soft_skills + " . " +
                self.model_config['form_clause'],
                self.model_config['resume_gen_temp']
            )

            response = self._parse_response(response)

            self.professional_experience_liminal[i]['all_relevant_responsibilities'] = response

        # select the most relevant responses
        for i in range(len(self.professional_experience_input)):
             response = complete_single_content(
                 client,
                 "Select the most relevant " + str(self.model_config['responsibility_count'][i]) + " " +
                 "responsibilities from this list: " + self._de_parse_response(self.professional_experience_liminal[i]['all_relevant_responsibilities']) + " . " 
                 "That cover as many of these areas as possible: " + self.gen_tech_skills + self.gen_tech_tools + self.gen_soft_skills + " . " +
                 self.model_config['form_clause'],
                 self.model_config['resume_gen_temp']
             )

             response = self._parse_response(response)

             self.professional_experience_liminal[i]['most_relevant_responsibilities'] = response

        # increase human readablity of responses
        for i in range(len(self.professional_experience_input)):
             response = complete_single_content(
                 client,
                 "Increase the readability of the elements in this list: " + self._de_parse_response(self.professional_experience_liminal[i]['most_relevant_responsibilities']) + " . " +
                 "Every statement should be clear and concise. " +
                 "Remove unnecessary punctuation. " +
                 "Ensure that all skills included here are preserved in the final output: " + self.gen_tech_skills + self.gen_tech_tools + self.gen_soft_skills + " . " +
                 "Additional extraneous skills may be removed if element is overly long. " +
                 self.model_config['form_clause'],
                 self.model_config['resume_gen_temp']
             )

             response = self._parse_response(response)

             self.professional_experience_liminal[i]['formatted_responsibilities'] = response

        # assign role titles for all employers
        for i in range(len(self.professional_experience_liminal)):
            if self.role_title_overrides[i] is not None:
                self.professional_experience_liminal[i]['role_title'] = self.role_title_overrides[i]
            else:
                role_title = self._generate_role_title(
                    client,
                    json.dumps(self.professional_experience_liminal[i]['formatted_responsibilities'])
                )
                self.professional_experience_liminal[i]['role_title'] = role_title

        # display role_title results to user
        string_output = "Generated role titles: \n"

        for i in range(len(self.professional_experience_liminal)):
             string_output += (
                 self.professional_experience_liminal[i]['employer'] + ": " +
                 self.professional_experience_liminal[i]['role_title'] + "\n"
             )

        log(string_output.rstrip("\n"))

        # assemble the final output resume
        for i in range(len(self.professional_experience_liminal)):
             # incorporate elements from input resume
             self.professional_experience_output.append({"employer": self.professional_experience_liminal[i]["employer"]})
             self.professional_experience_output[i]['role_title'] = self.professional_experience_liminal[i]['role_title']
             self.professional_experience_output[i]['employment_start'] = self.professional_experience_input[i]['employment_start']
             self.professional_experience_output[i]['employment_end'] = self.professional_experience_input[i]['employment_end']
             self.professional_experience_output[i]['responsibilities'] = self.professional_experience_liminal[i]['formatted_responsibilities']

        # inform user run was successful
        log('professional_experience output stored in GeneratedResume.professional_experience_output')


    def write_resume(self):
        # use still working field to determine display of pe0 employment_end
        log("writing resume")

        if self.doc_format['currently_employed']:
            self.professional_experience_output[0]['employment_end'] = ""

        # open and format doc
        resume_doc = Document()

        section = resume_doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # set document styles
        # header 1; will be used for "Professional Experience", "Education", etc.
        h1_style = resume_doc.styles['Heading 1']
        h1_font = h1_style.font
        h1_font.name = self.doc_format['h1']['font_name']
        h1_font.size = Pt(self.doc_format['h1']['font_size'])
        h1_font.color.rgb = RGBColor(
            self.doc_format['h1']['font_color'][0],
            self.doc_format['h1']['font_color'][1],
            self.doc_format['h1']['font_color'][2]
        )
        h1_paragraph_format = h1_style.paragraph_format
        h1_paragraph_format.line_spacing = self.doc_format['h1']['line_spacing']
        h1_paragraph_format.space_before = Pt(0)
        h1_paragraph_format.space_after = Pt(0)

        ## normal text; will be used for all other text
        normal_style = resume_doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.doc_format['normal']['font_name']
        normal_font.size = Pt(self.doc_format['normal']['font_size'])
        normal_font.color.rgb = RGBColor(
            self.doc_format['normal']['font_color'][0],
            self.doc_format['normal']['font_color'][1],
            self.doc_format['normal']['font_color'][2]
        )
        normal_paragraph_format = normal_style.paragraph_format
        normal_paragraph_format.line_spacing = self.doc_format['normal']['line_spacing']
        normal_paragraph_format.space_before = Pt(0)
        normal_paragraph_format.space_after = Pt(0)

        # sub_header formatting
        sub_header_font_size = normal_font.size
        sub_header_bold = True

        # insert either header image or text based on param
        if self.doc_format['use_image_header']:
            image_path = './data/assets/resume-header.png'
            resume_doc.add_picture(image_path, width=Inches(6.5))
        elif not self.doc_format['use_image_header']:
            resume_doc.add_heading(
                self.personal_info["first_name"].upper() + " " + self.personal_info["last_name"].upper(),
                level=1
            )
            resume_doc.add_paragraph(
                self.personal_info['email'] + "   |   " +
                self.personal_info['linkedin_url'] + "   |   " +
                self.personal_info['phone_number']
            )

        # horizontal bar
        # add_line(resume_doc)
        resume_doc.add_paragraph("")

        # professional experience
        # insert first section header or text header
        resume_doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
        self.add_line(resume_doc)

        # write text from model
        for i in range(len(self.professional_experience_output)):
            table = resume_doc.add_table(rows=1, cols=2)
            table.columns[0].width = Inches(4.5)
            table.columns[1].width = Inches(2.0)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = self.professional_experience_output[i]['role_title'] + ", " + self.professional_experience_output[i]['employer']
            hdr_cells[1].text = self.professional_experience_output[i]['employment_start'] + "-" + self.professional_experience_output[i][
                'employment_end']
            left_paragraph = hdr_cells[0].paragraphs[0]
            left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            right_paragraph = hdr_cells[1].paragraphs[0]
            right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = sub_header_bold
                        run.font.size = sub_header_font_size
                        run.font.name = 'Arial'
            for j in range(len(self.professional_experience_output[i]['responsibilities'])):
                resume_doc.add_paragraph(self.professional_experience_output[i]['responsibilities'][j],
                                         style='List Bullet')
            resume_doc.add_paragraph("")

        # other resume sections
        # education
        resume_doc.add_heading("EDUCATION", level=1)
        self.add_line(resume_doc)

        ## create table to house experience and date
        table = resume_doc.add_table(rows=1, cols=2)
        table.columns[0].width = Inches(4.5)
        table.columns[1].width = Inches(2.0)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = (
            self.education['degree'] + ' ' +
            self.education['major'] + ', ' +
            self.education['institution']
        )
        hdr_cells[1].text = (
            self.education['education_start'] + "-" +
            self.education['education_end']
        )
        left_paragraph = hdr_cells[0].paragraphs[0]
        left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        right_paragraph = hdr_cells[1].paragraphs[0]
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = sub_header_bold
                    run.font.size = sub_header_font_size
                    run.font.name = 'Arial'

        ## add minor
        resume_doc.add_paragraph(
            self.education['minor'],
            style='List Bullet'
        )
        resume_doc.add_paragraph("")

        ## military experience
        resume_doc.add_heading("MILITARY EXPERIENCE", level=1)
        self.add_line(resume_doc)

        ### create table to house experience and date
        table = resume_doc.add_table(rows=1, cols=2)
        table.columns[0].width = Inches(4.5)
        table.columns[1].width = Inches(2.0)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = (
            self.military_experience['role_title'] + ', ' +
            self.military_experience['branch']
        )
        hdr_cells[1].text = (
            self.military_experience['service_start'] + "-" +
            self.military_experience['service_end']
        )
        left_paragraph = hdr_cells[0].paragraphs[0]
        left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        right_paragraph = hdr_cells[1].paragraphs[0]
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = sub_header_bold
                    run.font.size = sub_header_font_size
                    run.font.name = self.doc_format['normal']['font_name']

        resume_doc.add_paragraph("")

        # hard skills
        resume_doc.add_heading("HARD SKILLS", level=1)
        self.add_line(resume_doc)

        for key, value in self.hard_skills.items():
            resume_doc.add_paragraph(
                f"{key}: {value}",
                style='List Bullet'
            )

        # build file output path and save doc
        resume_output_path = (
            self.env_vars['RESUME_OUTPUT_PATH'] +
            self.personal_info['first_name'].lower() + "-" +
            self.personal_info['last_name'].lower() + "-" +
            self.job_description['name_param'] +
            '-resume.docx'
        )
        resume_doc.save(resume_output_path)

        log('generated resume successfully written to: "' + resume_output_path + '"')


    def check_qualifications(self):
        """
        This function uses the OpenAI API to compare the qualifications of a job applicant to a job description.
        :return:
        """
        log("checking qualifications")
        # instantiate client
        client = OpenAI()

        # call API to assess the qualifications within the resume
        start_time = time.time()
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system",
                 "content": "The following json data describes the job skills of for Jeremy-Hatch: ```" + str(self.professional_experience_input) + "```"},
                {"role": "user", "content": "Follow these instructions: \
                    1. Compare the qualifications of Jeremy Hatch to the following job description:" + "```" + self.job_description['role_description'] + "```" +
                                            "2. Locate areas where Jeremy Hatch's qualifications may be lacking" +
                                            "3. Return only areas where qualifications may be lacking"
                 }
            ]
        )
        end_time = time.time()
        duration = end_time - start_time

        # write areas of improvement to areas_of_improvement.txt
        with open(self.env_vars['AREAS_OF_IMPROVEMENT_PATH'], "a") as file:
            file.write("## " + self.job_description['name_param'] + "\n")
            file.write(completion.choices[0].message.content)
            file.write("\n\n")

        # output response to user
        string_output = (
            f"API call duration:    {duration}\n" +
            f"input tokens:         {completion.usage.prompt_tokens}\n" +
            f"output tokens:        {completion.usage.completion_tokens}\n" +
            f"content:              {completion.choices[0].message.content}"
        )
        log(string_output)

# ------------------------------------------------------------------------------
# primary function
# ------------------------------------------------------------------------------

    def generate_resume(self):
        self.check_qualifications()
        self.generate_resume_content()
        self.write_resume()

# ------------------------------------------------------------------------------
# other external functions that are not part of the primary pipeline
# ------------------------------------------------------------------------------

    def pickle_resume(self):
        """
        pickle the resume object for later use
        """
        log("pickling resume object")
        resume_pickle_path = (
            self.env_vars['RESUME_OUTPUT_PATH'] +
            self.personal_info['first_name'].lower() + "-" +
            self.personal_info['last_name'].lower() + "-" +
            self.job_description['name_param'] +
            '-resume.pkl'
        )
        with open(resume_pickle_path, 'wb') as output:
            pickle.dump(self, output)

        log('resume object successfully pickled to: "' + resume_pickle_path + '"')

# ------------------------------------------------------------------------------
# end of generated_resume.py
# ------------------------------------------------------------------------------
