# external imports
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import dotenv_values
import json
import pathlib
import os
import yaml
# internal imports
from src.utils.single_content_completion import complete_single_content
from src.utils.logger import log

# ------------------------------------------------------------------------------
# load CoverLetter params and data
# ------------------------------------------------------------------------------

log('loading params and data')

env_vars = dotenv_values(".env")

with open('config/model_v1.25.yaml', 'r') as file:
    model_config: object = yaml.safe_load(file)

with open('config/doc_format.yaml', 'r') as file:
    doc_format = yaml.safe_load(file)

# ingest files if needed
resume_input_path = env_vars['RESUME_INPUT_PATH']
resume_input_path_sample = env_vars['RESUME_INPUT_PATH_SAMPLE']

# logical test to determine if file is present at RESUME_INPUT_PATH
try:
    if os.path.exists(resume_input_path):
        with open(resume_input_path, 'r') as file:
            resume_input = json.load(file)
    else:
        with open(resume_input_path_sample, 'r') as file:
            resume_input = json.load(file)
except Exception as e:
    print(f"Error reading files: {e}")

log('params and data loaded')

# ------------------------------------------------------------------------------
# define primary class
# ------------------------------------------------------------------------------

class GeneratedCoverLetter:
	def __init__(
		self,
		job_description,
		personal_info,
		resume,
		env_vars=env_vars,
		model_config=model_config,
		doc_format=doc_format
	):
		log("initializing GeneratedCoverLetter object")
		# input parameters
		self.resume = resume
		self.env_vars = env_vars
		self.doc_format = doc_format
		self.job_description = job_description
		self.model_config = model_config
		self.personal_info = personal_info
		self.resume = resume
		# generated content to be defined via methods
		self.cover_letter_text = None
		self.cover_letter_output_path = None
		log("GeneratedCoverLetter object initialized")


# ------------------------------------------------------------------------------
# sub methods to main cover letter generator method
# -----------------------------------------------------------------------------

	def generate_cover_letter_content(self):
		"""generates a cover letter using the OpenAI API"""
		log("generating cover letter content")

		# read in user generated cover letter content if any
		cl_content_path = self.env_vars['COVER_LETTER_CONTENT_PATH'] + self.job_description['company_name'] + '.txt'

		if pathlib.Path(cl_content_path).exists():
			with open(self.env_vars['COVER_LETTER_CONTENT_PATH'], 'r', encoding='utf-8') as f:
				cl_content = f.read()
		else:
			cl_content = None

		# retrieve information about company
		company_info = complete_single_content(
			"Give me a summary about this company and tell me about its " + \
			"values: " + self.job_description['company_name']
		)

		# generate body of cover letter
		if cl_content is not None:
			self.cover_letter_text = complete_single_content(
				"An individual has this experience: " + str(self.resume) + ". " +
				"They are applying for this job: " + self.job_description['role_description'] + ". " +
				"at this company: " + company_info + " ." +
				"Write them a cover letter following these instructions: " +
				"1. Incorporate this content ```" + cl_content + "``` in a logical fashion; perform grammatical and structural improvements to this content if need be. " +
				"2. Output should contain the body of the cover letter and not the address or salutation." +
				"3. Focus on why the applicant is interested in the the key areas of the role and the company." +
				"4. Do not use the word 'excitement' or any synonym in the first sentence." +
				"5. Do not overly emphasize the applicant's experience, this should not be a duplicate of their resume." +
				"6. Do not include a complimentary closing. " +
				"7. Limit the output to 4 paragraphs."
			)
		else:
			self.cover_letter_text = complete_single_content(
				"An individual has this experience: " + str(self.resume) + ". " +
				"They are applying for this job: " + self.job_description['role_description'] + ". " +
				"at this company: " + company_info + " ." +
				"Write them a cover letter following these instructions: " +
				"1. Output should contain the body of the cover letter and not the address or salutation." +
				"2. Focus on why the applicant is interested in the the key areas of the role and the company." +
				"3. Do not use the word 'excitement' or any synonym in the first sentence." +
				"4. Do not overly emphasize the applicant's experience, this should not be a duplicate of their resume." +
				"5. Do not include a complimentary closing. " +
				"6. Limit the output to 4 paragraphs."
			)

		log("cover letter content generated")


	def write_cover_letter(self):
		"""
		takes the cover_letter_text string from the cover_letter_gen function and
		writes it to a .docx file at the specified path
		"""
		log("writing cover letter to docx file")

		# open and format doc
		cover_letter_doc = Document()

		section = cover_letter_doc.sections[0]
		section.top_margin = Inches(1)
		section.bottom_margin = Inches(1)
		section.left_margin = Inches(1)
		section.right_margin = Inches(1)

		# set document styles
		normal_style = cover_letter_doc.styles['Normal']
		normal_font = normal_style.font
		normal_font.name = self.doc_format['cover_letter']['font_name']
		normal_font.size = Pt(self.doc_format['cover_letter']['font_size'])
		normal_font.color.rgb = RGBColor(
			self.doc_format['cover_letter']['font_color'][0],
			self.doc_format['cover_letter']['font_color'][1],
			self.doc_format['cover_letter']['font_color'][2]
		)
		normal_paragraph_format = normal_style.paragraph_format
		normal_paragraph_format.line_spacing = self.doc_format['cover_letter']['line_spacing']
		normal_paragraph_format.space_before = Pt(0)
		normal_paragraph_format.space_after = Pt(0)

		# insert either header image or text based on param
		if self.doc_format['cover_letter_user_image_header']:
			image_path = './data/assets/resume-header.png'
			cover_letter_doc.add_picture(image_path, width=Inches(6.5))
		else:
			heading = cover_letter_doc.add_heading(
                self.personal_info["first_name"].upper() + " " + self.personal_info["last_name"].upper(),
                level=1
            )
			heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
			contact_info = cover_letter_doc.add_paragraph(
                self.personal_info['email'] + "   |   " +
                self.personal_info['linkedin_url'] + "   |   " +
                self.personal_info['phone_number']
            )
			contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

		# adding white space between header and cover letter text
		cover_letter_doc.add_paragraph()
		cover_letter_doc.add_paragraph()

		# add opening
		cover_letter_doc.add_paragraph("Dear Hiring Manager, \n")

		# add cover letter text to doc
		cover_letter_doc.add_paragraph(self.cover_letter_text)

		# adding complimentary closing
		cover_letter_doc.add_paragraph()
		cover_letter_doc.add_paragraph(
			"Looking forward to meeting you, \n" +
			self.personal_info["first_name"] + " " +
			self.personal_info["last_name"]
		)

		# build cover letter output path and save
		self.cover_letter_output_path = (
			self.env_vars['COVER_LETTER_OUTPUT_PATH'] +
			self.personal_info['first_name'].lower() + "-" +
			self.personal_info['last_name'].lower() + "-" +
			self.job_description['name_param'] +
			"-cover-letter.docx"
		)

		cover_letter_doc.save(self.cover_letter_output_path)

		log("cover letter saved to " + self.cover_letter_output_path)

# ------------------------------------------------------------------------------
# primary cover letter generation method
# -----------------------------------------------------------------------------_

	def generate_cover_letter(self):
		"""
		master cover letter generation function that generates all content and
		write the file to the generated path
		"""
		log("initiating cover letter generation pipeline")
		self.generate_cover_letter_content()
		self.write_cover_letter()

# ------------------------------------------------------------------------------
# end of generate_cover_letter.py
# -----------------------------------------------------------------------------_
